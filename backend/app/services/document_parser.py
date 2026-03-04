"""Document parser service — extract structured rules from policy docs, contracts, and audit reports.

Supports .docx (python-docx) and .pdf (pdfplumber) formats.
Uses Claude AI for intelligent rule extraction when available.
"""

from __future__ import annotations

import json
import logging
import re
import uuid
from pathlib import Path
from typing import Any

from sqlalchemy.orm import Session

from app.models.config import (
    ExtractionStatus,
    PolicyDocument,
    PolicyRule,
    PolicyRuleStatus,
)
from app.services.ai_service import ai_service

logger = logging.getLogger(__name__)

# ── Document type constants ──────────────────────────────────────────────

DOC_TYPE_POLICY = "policy"
DOC_TYPE_CONTRACT = "contract"
DOC_TYPE_AUDIT = "audit_report"
DOC_TYPE_MATCH_REPORT = "match_report"


# ── Text extraction helpers ──────────────────────────────────────────────


def extract_text_from_docx(file_path: str) -> str:
    """Extract raw text from a .docx file."""
    import docx

    doc = docx.Document(file_path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # Also extract table content
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines)


def extract_text_from_pdf(file_path: str) -> str:
    """Extract raw text from a .pdf file."""
    import pdfplumber

    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

            # Also extract table data
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    cells = [str(cell).strip() for cell in row if cell]
                    if cells:
                        text_parts.append(" | ".join(cells))

    return "\n".join(text_parts)


def extract_text(file_path: str) -> str:
    """Extract text from a document, auto-detecting format."""
    path = Path(file_path)
    ext = path.suffix.lower()
    if ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".pdf":
        return extract_text_from_pdf(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


# ── AI-powered rule extraction ───────────────────────────────────────────

POLICY_EXTRACTION_PROMPT = """You are an AP (Accounts Payable) policy analyst. Extract structured business rules from the following policy document.

For each rule, provide:
- rule_type: one of [approval_threshold, matching_requirement, exception_handling, payment_terms, vendor_management, audit_control, kpi_target, invoice_validation, duplicate_prevention, tax_treatment, escalation]
- source_section: which section of the document this comes from
- source_text: the original text (max 200 chars)
- conditions: JSON object describing when this rule applies (e.g., {"amount_min": 10001, "amount_max": 50000})
- action: JSON object describing what should happen (e.g., {"approver": "AP Manager", "auto_approve": false})
- confidence: your confidence this is a real business rule (0.0-1.0)

Return a JSON array of rules. Extract ALL actionable rules — threshold values, SLA targets, required fields, escalation paths, approval levels, matching tolerances, etc.

DOCUMENT TEXT:
"""

CONTRACT_EXTRACTION_PROMPT = """You are a contract analyst for Accounts Payable. Extract structured terms from this supplier contract.

For each term, provide:
- rule_type: one of [payment_terms, price_tolerance, surcharge_allowance, volume_discount, delivery_terms, quality_terms, dispute_resolution, penalty_clause, invoicing_requirement]
- source_section: which section this comes from
- source_text: the original text (max 200 chars)
- conditions: JSON with conditions (e.g., {"supplier": "SteelCore", "order_min": 50000})
- action: JSON with the term details (e.g., {"discount_pct": 2.0, "payment_days": 45})
- confidence: 0.0-1.0

Return a JSON array. Extract ALL commercial terms — payment periods, discounts, tolerances, surcharges, penalties, SLAs, etc.

CONTRACT TEXT:
"""

AUDIT_EXTRACTION_PROMPT = """You are an internal audit analyst. Extract structured findings from this AP audit report.

For each finding, provide:
- rule_type: one of [audit_finding, control_gap, process_improvement, kpi_deviation, vendor_issue, compliance_risk]
- source_section: which section/finding number
- source_text: brief description (max 200 chars)
- conditions: JSON with finding details (e.g., {"severity": "high", "finding_ref": "FINDING-1"})
- action: JSON with recommendation (e.g., {"recommendation": "...", "due_date": "...", "status": "open"})
- confidence: 0.0-1.0

Return a JSON array. Extract ALL findings, recommendations, KPI gaps, and management responses.

AUDIT REPORT TEXT:
"""


def _get_extraction_prompt(doc_type: str) -> str:
    """Get the appropriate extraction prompt for a document type."""
    prompts = {
        DOC_TYPE_POLICY: POLICY_EXTRACTION_PROMPT,
        DOC_TYPE_CONTRACT: CONTRACT_EXTRACTION_PROMPT,
        DOC_TYPE_AUDIT: AUDIT_EXTRACTION_PROMPT,
        DOC_TYPE_MATCH_REPORT: AUDIT_EXTRACTION_PROMPT,
    }
    return prompts.get(doc_type, POLICY_EXTRACTION_PROMPT)


def extract_rules_with_ai(text: str, doc_type: str) -> list[dict[str, Any]]:
    """Use Claude AI to extract structured rules from document text."""
    if not ai_service.available:
        logger.warning("AI service unavailable — using fallback rule extraction")
        return _extract_rules_fallback(text, doc_type)

    prompt = _get_extraction_prompt(doc_type)
    response = ai_service.call_claude(
        system_prompt="You are a precise rule extraction engine. Return ONLY valid JSON arrays.",
        user_message=prompt + text[:12000],  # Limit to avoid token limits
        max_tokens=4096,
    )

    if not response:
        logger.warning("AI extraction returned empty — using fallback")
        return _extract_rules_fallback(text, doc_type)

    # Parse JSON from response
    try:
        # Try to find JSON array in response
        match = re.search(r"\[.*\]", response, re.DOTALL)
        if match:
            rules = json.loads(match.group())
            if isinstance(rules, list):
                return rules
    except json.JSONDecodeError:
        pass

    logger.warning("Failed to parse AI extraction response — using fallback")
    return _extract_rules_fallback(text, doc_type)


# ── Fallback rule extraction (regex/pattern-based) ──────────────────────


def _extract_rules_fallback(text: str, doc_type: str) -> list[dict[str, Any]]:
    """Pattern-based fallback when AI is unavailable."""
    rules = []

    if doc_type == DOC_TYPE_POLICY:
        rules.extend(_extract_policy_rules_fallback(text))
    elif doc_type == DOC_TYPE_CONTRACT:
        rules.extend(_extract_contract_rules_fallback(text))
    elif doc_type in (DOC_TYPE_AUDIT, DOC_TYPE_MATCH_REPORT):
        rules.extend(_extract_audit_rules_fallback(text))

    return rules


def _extract_policy_rules_fallback(text: str) -> list[dict[str, Any]]:
    """Extract key policy rules via pattern matching."""
    rules = []

    # Approval thresholds
    threshold_patterns = [
        (r"\$0\s*[–-]\s*\$?(\d[\d,]*)", r"Auto-approved"),
        (r"\$(\d[\d,]*)\s*[–-]\s*\$?(\d[\d,]*)\s*:?\s*(.*?)(?:\n|$)", None),
        (r"Above\s*\$(\d[\d,]*)\s*:?\s*(.*?)(?:\n|$)", None),
    ]
    for pattern, _ in threshold_patterns:
        for m in re.finditer(pattern, text, re.IGNORECASE):
            rules.append(
                {
                    "rule_type": "approval_threshold",
                    "source_section": "Section 5 — Approval Authority Matrix",
                    "source_text": m.group(0)[:200],
                    "conditions": {"threshold_text": m.group(0)[:100]},
                    "action": {"approval_required": True},
                    "confidence": 0.8,
                }
            )

    # Matching requirements
    if "3-way match" in text.lower():
        rules.append(
            {
                "rule_type": "matching_requirement",
                "source_section": "Section 4.1 — 3-Way Match",
                "source_text": "All goods invoices require 3-way match: Invoice vs PO vs GRN",
                "conditions": {"invoice_type": "goods"},
                "action": {"match_type": "3-way", "fields": ["invoice", "po", "grn"]},
                "confidence": 0.95,
            }
        )

    if "2-way match" in text.lower():
        rules.append(
            {
                "rule_type": "matching_requirement",
                "source_section": "Section 4.2 — 2-Way Match",
                "source_text": "Service invoices with valid contract reference may use 2-way match",
                "conditions": {"invoice_type": "services", "has_contract": True},
                "action": {"match_type": "2-way", "fields": ["invoice", "po"]},
                "confidence": 0.9,
            }
        )

    # Exception SLA targets
    sla_pattern = r"(Price Variance|Quantity Mismatch|Missing PO|Duplicate Invoice|Unknown Supplier|Description Mismatch|Incorrect Tax|Above Threshold)\s*[│|]\s*(\d+)"
    for m in re.finditer(sla_pattern, text):
        rules.append(
            {
                "rule_type": "exception_handling",
                "source_section": "Section 6.1 — Exception Types",
                "source_text": f"{m.group(1)}: {m.group(2)} hours target resolution",
                "conditions": {"exception_type": m.group(1).lower().replace(" ", "_")},
                "action": {"target_hours": int(m.group(2))},
                "confidence": 0.9,
            }
        )

    # KPI targets
    kpi_pattern = r"(Touchless Rate|Invoice Cycle Time|Exception Rate|On-Time Payment Rate|Aging >30 days)\s*[│|]\s*([\d<>%.\s]+days?)"
    for m in re.finditer(kpi_pattern, text):
        rules.append(
            {
                "rule_type": "kpi_target",
                "source_section": "Section 10 — KPI Targets",
                "source_text": f"{m.group(1)}: target {m.group(2).strip()}",
                "conditions": {"metric": m.group(1)},
                "action": {"target": m.group(2).strip()},
                "confidence": 0.85,
            }
        )

    # Duplicate prevention
    if "duplicate" in text.lower():
        rules.append(
            {
                "rule_type": "duplicate_prevention",
                "source_section": "Section 3.3 — Duplicate Invoice Prevention",
                "source_text": "System must check supplier + invoice number combination before processing",
                "conditions": {"check_fields": ["supplier_id", "invoice_number"]},
                "action": {"auto_reject": True, "notify_supplier": True},
                "confidence": 0.95,
            }
        )

    return rules


def _extract_contract_rules_fallback(text: str) -> list[dict[str, Any]]:
    """Extract key contract terms via pattern matching."""
    rules = []

    # Extract supplier name
    supplier_match = re.search(r"Supplier:\s*(.*?)(?:\n|$)", text)
    supplier = supplier_match.group(1).strip() if supplier_match else "Unknown"

    # Payment terms
    payment_match = re.search(r"Payment Terms:\s*(Net\s*\d+)", text, re.IGNORECASE)
    if payment_match:
        days = re.search(r"\d+", payment_match.group(1))
        rules.append(
            {
                "rule_type": "payment_terms",
                "source_section": "Commercial Terms",
                "source_text": f"{supplier}: {payment_match.group(1)}",
                "conditions": {"supplier": supplier},
                "action": {"payment_days": int(days.group()) if days else 30},
                "confidence": 0.95,
            }
        )

    # Price tolerance
    tolerance_match = re.search(r"Price Tolerance:\s*(.*?)(?:\n|$)", text, re.IGNORECASE)
    if tolerance_match:
        rules.append(
            {
                "rule_type": "price_tolerance",
                "source_section": "Invoicing Rules",
                "source_text": f"{supplier}: {tolerance_match.group(1).strip()}",
                "conditions": {"supplier": supplier},
                "action": {"tolerance": tolerance_match.group(1).strip()},
                "confidence": 0.9,
            }
        )

    # Surcharges
    surcharge_match = re.search(r"Maximum:\s*([\d.]+%.*?)(?:\n|$)", text, re.IGNORECASE)
    if surcharge_match:
        type_match = re.search(r"Type:\s*(.*?)(?:\n|$)", text)
        rules.append(
            {
                "rule_type": "surcharge_allowance",
                "source_section": "Permitted Surcharges",
                "source_text": f"{supplier}: {type_match.group(1).strip() if type_match else 'surcharge'} up to {surcharge_match.group(1).strip()}",
                "conditions": {"supplier": supplier},
                "action": {"max_surcharge": surcharge_match.group(1).strip()},
                "confidence": 0.9,
            }
        )

    # Volume discounts
    for m in re.finditer(r"Orders?\s*>\s*\$?([\d,]+)\s*:\s*([\d.]+%)\s*discount", text, re.IGNORECASE):
        rules.append(
            {
                "rule_type": "volume_discount",
                "source_section": "Volume Discounts",
                "source_text": f"{supplier}: {m.group(2)} discount for orders > ${m.group(1)}",
                "conditions": {"supplier": supplier, "order_min": float(m.group(1).replace(",", ""))},
                "action": {"discount_pct": float(m.group(2).replace("%", ""))},
                "confidence": 0.9,
            }
        )

    # Late delivery penalty
    penalty_match = re.search(r"Late Delivery Penalty:\s*(.*?)(?:\n|$)", text, re.IGNORECASE)
    if penalty_match:
        rules.append(
            {
                "rule_type": "penalty_clause",
                "source_section": "Delivery Terms",
                "source_text": f"{supplier}: {penalty_match.group(1).strip()}",
                "conditions": {"supplier": supplier, "trigger": "late_delivery"},
                "action": {"penalty": penalty_match.group(1).strip()},
                "confidence": 0.85,
            }
        )

    return rules


def _extract_audit_rules_fallback(text: str) -> list[dict[str, Any]]:
    """Extract audit findings via pattern matching."""
    rules = []

    # Extract findings
    finding_pattern = (
        r"FINDING\s*(\d+)\s*[—-]\s*(HIGH|MEDIUM|LOW)\s+SEVERITY\s*\n.*?Category:\s*(.*?)\n.*?Title:\s*(.*?)\n"
    )
    for m in re.finditer(finding_pattern, text, re.IGNORECASE | re.DOTALL):
        rules.append(
            {
                "rule_type": "audit_finding",
                "source_section": f"Finding {m.group(1)}",
                "source_text": f"[{m.group(2)}] {m.group(4).strip()}",
                "conditions": {
                    "severity": m.group(2).lower(),
                    "category": m.group(3).strip(),
                    "finding_ref": f"FINDING-{m.group(1)}",
                },
                "action": {"status": "open"},
                "confidence": 0.9,
            }
        )

    return rules


# ── Main parse + store function ──────────────────────────────────────────


def parse_and_store_document(
    db: Session,
    file_path: str,
    filename: str,
    doc_type: str,
    uploaded_by: uuid.UUID | None = None,
    use_ai: bool = True,
) -> PolicyDocument:
    """Parse a document, extract rules, and store in the database.

    Returns the created PolicyDocument with linked PolicyRules.
    """
    # Create the document record
    doc = PolicyDocument(
        filename=filename,
        file_path=file_path,
        document_type=doc_type,
        uploaded_by=uploaded_by,
        extraction_status=ExtractionStatus.processing,
    )
    db.add(doc)
    db.flush()

    try:
        # Extract text
        text = extract_text(file_path)
        logger.info("Extracted %d chars from %s", len(text), filename)

        # Extract rules
        if use_ai and ai_service.available:
            raw_rules = extract_rules_with_ai(text, doc_type)
        else:
            raw_rules = _extract_rules_fallback(text, doc_type)

        logger.info("Extracted %d rules from %s", len(raw_rules), filename)

        # Store rules
        for rule_data in raw_rules:
            rule = PolicyRule(
                policy_document_id=doc.id,
                rule_type=rule_data.get("rule_type", "unknown"),
                source_text=rule_data.get("source_text", "")[:500],
                conditions=rule_data.get("conditions"),
                action=rule_data.get("action"),
                confidence=rule_data.get("confidence", 0.5),
                status=PolicyRuleStatus.pending,
            )
            db.add(rule)

        doc.extraction_status = ExtractionStatus.completed
        doc.extracted_rules_count = len(raw_rules)
        db.commit()

        logger.info("Successfully parsed %s: %d rules stored", filename, len(raw_rules))
        return doc

    except Exception as e:
        doc.extraction_status = ExtractionStatus.failed
        db.commit()
        logger.error("Failed to parse %s: %s", filename, e)
        raise


def parse_all_ap_inputs(db: Session, ap_inputs_dir: str, use_ai: bool = True) -> dict[str, Any]:
    """Parse all documents from the AP_Inputs directory.

    Returns a summary of parsed documents and extracted rules.
    """
    base = Path(ap_inputs_dir)
    results = {"documents": [], "total_rules": 0}

    # Parse AP Policy
    policy_path = base / "AP_Policy.docx"
    if policy_path.exists():
        doc = parse_and_store_document(db, str(policy_path), "AP_Policy.docx", DOC_TYPE_POLICY, use_ai=use_ai)
        results["documents"].append(
            {
                "filename": doc.filename,
                "type": DOC_TYPE_POLICY,
                "rules_extracted": doc.extracted_rules_count,
            }
        )
        results["total_rules"] += doc.extracted_rules_count

    # Parse 3-Way Match Report
    match_report = base / "AP_3Way_Match_Report.docx"
    if match_report.exists():
        doc = parse_and_store_document(
            db, str(match_report), "AP_3Way_Match_Report.docx", DOC_TYPE_MATCH_REPORT, use_ai=use_ai
        )
        results["documents"].append(
            {
                "filename": doc.filename,
                "type": DOC_TYPE_MATCH_REPORT,
                "rules_extracted": doc.extracted_rules_count,
            }
        )
        results["total_rules"] += doc.extracted_rules_count

    # Parse Audit Findings
    audit_path = base / "Audit_Findings_2024.pdf"
    if audit_path.exists():
        doc = parse_and_store_document(db, str(audit_path), "Audit_Findings_2024.pdf", DOC_TYPE_AUDIT, use_ai=use_ai)
        results["documents"].append(
            {
                "filename": doc.filename,
                "type": DOC_TYPE_AUDIT,
                "rules_extracted": doc.extracted_rules_count,
            }
        )
        results["total_rules"] += doc.extracted_rules_count

    # Parse Supplier Contracts
    contracts_dir = base / "Supplier_Contracts"
    if contracts_dir.exists():
        for contract_path in sorted(contracts_dir.glob("*.docx")):
            doc = parse_and_store_document(db, str(contract_path), contract_path.name, DOC_TYPE_CONTRACT, use_ai=use_ai)
            results["documents"].append(
                {
                    "filename": doc.filename,
                    "type": DOC_TYPE_CONTRACT,
                    "rules_extracted": doc.extracted_rules_count,
                }
            )
            results["total_rules"] += doc.extracted_rules_count

    return results
